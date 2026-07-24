import pdfplumber
# pyrefly: ignore [missing-import]
import PyPDF2
from docx import Document

def extract_pdf_text_and_pages(file_path) -> tuple:
    """
    Extracts text contents and counts pages from PDF files.
    Tries pdfplumber first, falling back to PyPDF2 upon failure or empty results.
    """
    text = []
    pages_count = 0
    
    # Primary parsing: pdfplumber
    try:
        with pdfplumber.open(file_path) as pdf:
            pages_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        extracted = "\n".join(text).strip()
        if extracted:
            return extracted, pages_count
    except Exception:
        # Fail silently to trigger fallback
        pass

    # Fallback parsing: PyPDF2
    try:
        # Re-initialize lists
        text = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            pages_count = len(reader.pages)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return "\n".join(text).strip(), max(pages_count, 1)
    except Exception:
        return "", 1


def extract_docx_text_and_pages(file_path) -> tuple:
    """
    Extracts text contents and counts pages from Word (.docx) files.
    Calculates page counts by inspecting rendering break tags in the doc XML structure.
    """
    try:
        doc = Document(file_path)
        text = []
        
        # Extract from normal paragraphs
        for para in doc.paragraphs:
            if para.text:
                text.append(para.text)
                
        # Extract from tables structures
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text.append(cell.text)
                        
        # Estimate page count by counting standard page breaks in document XML
        doc_xml = doc.element.xml
        xml_page_breaks = doc_xml.count('w:br w:type="page"')
        xml_rendered_breaks = doc_xml.count('w:lastRenderedPageBreak')
        
        pages_count = xml_page_breaks + xml_rendered_breaks + 1
        
        return "\n".join(text).strip(), max(pages_count, 1)
    except Exception:
        return "", 1


def parse_resume_document(file_path, file_type: str) -> tuple:
    """
    Router utility to parse document text and count pages based on file extension type.
    """
    ft = file_type.lower().strip('.')
    if ft == 'pdf':
        return extract_pdf_text_and_pages(file_path)
    elif ft in ['doc', 'docx']:
        return extract_docx_text_and_pages(file_path)
    return "", 1


class JobMatchGenerator:
    SYSTEM_PROMPT = """
You are "PrepAI Job Description Matcher" — an expert AI recruiter and talent analyst.
Your task is to analyze a candidate's resume text against a target Job Description (JD) text, perform a professional skill-gap analysis, and provide clear recommendations.

RULES:
1. Extract candidate's skills strictly from the resume. Do not assume skills not mentioned or implied by actual accomplishments.
2. Extract required skills from the job description.
3. Compare the candidate's skills against the job description requirements to identify:
   - Matching Skills (skills in the resume that match the JD requirements)
   - Missing Skills (critical/important skills in the JD that are missing or weak in the resume)
4. Calculate a match_score (0 to 100) representing how well the candidate's profile aligns with the JD. Be realistic:
   - 90-100: Perfect or near-perfect match of all required tech stack and experience.
   - 70-89: Good match, possesses most core skills, minor gaps in secondary requirements.
   - 50-69: Moderate match, has some matching core skills but misses several key requirements.
   - Below 50: Weak match, fits very few requirements.
5. Provide a professional list of strengths (why they fit) and a constructive gap analysis (where they fall short).
6. Provide actionable recommendations on how to update their resume or prepare for the interview to address the gaps.
7. Respond ONLY in the requested JSON format. No markdown block wrapper, no preamble, no tail.
"""

    @classmethod
    def build_user_prompt(cls, resume_text: str, job_description: str) -> str:
        return f"""
RESUME TEXT:
\"\"\"
{resume_text}
\"\"\"

JOB DESCRIPTION (JD):
\"\"\"
{job_description}
\"\"\"

TASK:
Perform a professional match analysis. Output a JSON object matching this schema:
{{
  "match_score": 85,
  "detected_candidate_skills": ["Python", "Django", "PostgreSQL", "React", "REST APIs"],
  "required_job_skills": ["Python", "Django", "Kubernetes", "AWS", "gRPC"],
  "matching_skills": ["Python", "Django"],
  "missing_skills": ["Kubernetes", "AWS", "gRPC"],
  "strengths": [
    "Strong backend engineering foundation with Django and Python.",
    "Good experience in REST APIs and database modeling."
  ],
  "gap_analysis": [
    "Missing container orchestration experience (Kubernetes).",
    "No explicit AWS cloud hosting experience mentioned in the resume."
  ],
  "actionable_recommendations": [
    "Add any project work involving Kubernetes or Docker containers to highlight your deployment skills.",
    "If you have used AWS, explicitly mention services like EC2, S3, or ECS under your professional experience."
  ]
}}
"""

